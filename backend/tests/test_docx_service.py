import unittest
import asyncio
from io import BytesIO
import base64
from unittest.mock import patch, AsyncMock, MagicMock
from docx import Document
from services.docx_service import DOCXService  # Now these imports should resolve

# In order to run this, put documentParsing folder into the backend folder and run:
# python3 -m unittest tests.test_docx_service

class TestDOCXService(unittest.TestCase):

    # Sets up a mock document with just hello world in it, b64
    def setUp(self):
        self.service = DOCXService()

        # Create a simple in-memory DOCX file
        doc = Document()
        doc.add_paragraph("Hello, world!")
        fake_docx = BytesIO()
        doc.save(fake_docx)
        self.base64_docx = base64.b64encode(fake_docx.getvalue()).decode("utf-8")

    # Tests extract text from docx
    @patch("documentParsing.parsers.docx_parser.parse_docx")
    def test_extract_text_from_docx(self, mock_parse_docx):
        mock_parse_docx.return_value = [{"text": "Hello, world!"}]

        async def run_test():
            extracted_text, text_map = await self.service.extract_text_from_docx(self.base64_docx)
            self.assertEqual(extracted_text, "Hello, world!")
            self.assertIsInstance(text_map, list)

        asyncio.run(run_test())

    # Tests write text to memory
    @patch("documentParsing.overlay.docx_overlay.overlay_docx")
    def test_write_text_to_memory(self, mock_overlay_docx):
       # Just return the original document to simulate overlay
       def fake_overlay(doc, text_map, translated_texts, image_map):
          return doc  # No modification, just return doc to test flow

       mock_overlay_docx.side_effect = fake_overlay

       translated = "This is translated text\nAnother line"
       text_map = [{"text": "Hello, world!", "type": "paragraph", "index": 0}, {"text": "Line 2", "type": "paragraph", "index": 0}]

       try:
          result_bytes = self.service.write_text_to_memory(translated, self.base64_docx, text_map)
          # Validate it's actually a .docx
          result_doc = Document(BytesIO(result_bytes))
          self.assertIsInstance(result_bytes, bytes)
          self.assertTrue(len(result_doc.paragraphs) >= 1)
       except Exception as e:
          self.fail(f"write_text_to_memory raised an exception: {e}")

    # Tests invalid input and ensures an error is raised
    def test_extract_text_invalid_input(self):
        async def run_test():
            with self.assertRaises(Exception) as context:
                await self.service.extract_text_from_docx(None)
            self.assertIn("No DOCX file provided", str(context.exception))

        asyncio.run(run_test())

if __name__ == "__main__":
    unittest.main()
